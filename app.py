import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import numpy as np
import docx
import fitz  # PyMuPDF
from PIL import Image
import pytesseract
import requests
import json
import re
import time
import io
import sys
from datetime import datetime
import plotly.express as px
import plotly.graph_objects as go

# --- Configuration ---
st.set_page_config(
    page_title="Advanced Data Analyst Agent",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Initialize session state
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'current_data' not in st.session_state:
    st.session_state.current_data = None
if 'data_info' not in st.session_state:
    st.session_state.data_info = {}
if 'current_file_name' not in st.session_state:
    st.session_state.current_file_name = None
if 'current_chat_id' not in st.session_state:
    st.session_state.current_chat_id = None
if 'data_summary_cache' not in st.session_state:
    st.session_state.data_summary_cache = {}

# --- API Configuration ---
# --- API Configuration ---
import streamlit as st

try:
    api_key = st.secrets["TOGETHER_API_KEY"]
except KeyError:
    st.error("⚙️ Server configuration in progress. Please refresh in a few seconds.")
    st.stop()

LLAMA_MODEL_ID = "meta-llama/Llama-4-Maverick-17B-128E-Instruct-FP8"

# --- OPTIMIZATION 1: Smart Data Summarization with Caching ---
@st.cache_data(ttl=600)
def get_compact_data_summary(df, file_name, max_rows=3):
    """Generate minimal, cached data summary for faster prompts"""
    numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
    categorical_cols = df.select_dtypes(include=['object']).columns.tolist()
    
    summary = {
        'shape': df.shape,
        'columns': list(df.columns)[:12],  # Limit to 12 columns
        'numeric_cols': numeric_cols[:5],
        'categorical_cols': categorical_cols[:5],
        'sample': df.head(max_rows).to_dict('records'),
        'has_nulls': df.isnull().sum().sum() > 0
    }
    return summary

# --- Enhanced File Parser ---
@st.cache_data
def parse_file(uploaded_file):
    """Parse uploaded file and return data with metadata"""
    try:
        suffix = uploaded_file.name.split(".")[-1].lower()
        file_size = uploaded_file.size
        
        if suffix == "csv":
            try:
                df = pd.read_csv(uploaded_file, encoding='utf-8')
            except UnicodeDecodeError:
                df = pd.read_csv(uploaded_file, encoding='latin-1')
            
            metadata = {
                'type': 'dataframe',
                'shape': df.shape,
                'columns': list(df.columns),
                'dtypes': df.dtypes.to_dict(),
                'file_size': file_size,
                'missing_values': df.isnull().sum().to_dict()
            }
            return df, "dataframe", metadata
            
        elif suffix == "xlsx":
            df = pd.read_excel(uploaded_file)
            metadata = {
                'type': 'dataframe',
                'shape': df.shape,
                'columns': list(df.columns),
                'dtypes': df.dtypes.to_dict(),
                'file_size': file_size,
                'missing_values': df.isnull().sum().to_dict()
            }
            return df, "dataframe", metadata
            
        elif suffix == "txt":
            content = uploaded_file.read().decode('utf-8')
            metadata = {
                'type': 'text',
                'length': len(content),
                'word_count': len(content.split()),
                'file_size': file_size
            }
            return content, "text", metadata
            
        elif suffix == "docx":
            doc = docx.Document(uploaded_file)
            content = "\n".join([p.text for p in doc.paragraphs])
            metadata = {
                'type': 'text',
                'paragraphs': len(doc.paragraphs),
                'length': len(content),
                'word_count': len(content.split()),
                'file_size': file_size
            }
            return content, "text", metadata
            
        elif suffix == "pdf":
            doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
            content = "\n".join([page.get_text() for page in doc])
            metadata = {
                'type': 'text',
                'pages': len(doc),
                'length': len(content),
                'word_count': len(content.split()),
                'file_size': file_size
            }
            return content, "text", metadata
            
        elif suffix in ["png", "jpg", "jpeg"]:
            image = Image.open(uploaded_file)
            content = pytesseract.image_to_string(image)
            metadata = {
                'type': 'text',
                'image_size': image.size,
                'extracted_text_length': len(content),
                'file_size': file_size
            }
            return content, "text", metadata
            
        else:
            return "Unsupported file type", None, {}
            
    except Exception as e:
        st.error(f"Error parsing file: {str(e)}")
        return None, None, {}

# --- OPTIMIZATION 2: Streaming LLM Response ---
def query_llama_streaming(prompt, max_tokens=800):
    """Streaming LLM for real-time response display"""
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }

    payload = {
        "model": LLAMA_MODEL_ID,
        "prompt": prompt,
        "max_tokens": max_tokens,
        "temperature": 0.7,
        "stream": True
    }

    try:
        response = requests.post(
            "https://api.together.xyz/v1/completions",
            headers=headers,
            json=payload,
            timeout=25,
            stream=True
        )
        
        if response.status_code == 200:
            full_text = ""
            placeholder = st.empty()
            
            for line in response.iter_lines():
                if line:
                    try:
                        line_text = line.decode('utf-8')
                        if line_text.startswith('data: '):
                            json_str = line_text[6:]
                            if json_str.strip() == '[DONE]':
                                break
                            data = json.loads(json_str)
                            if 'choices' in data and len(data['choices']) > 0:
                                token = data['choices'][0].get('text', '')
                                full_text += token
                                placeholder.markdown(full_text + "▌")
                    except:
                        continue
            
            placeholder.markdown(full_text)
            return full_text.strip()
        else:
            result = response.json()
            return f"❌ API Error: {result.get('error', 'Unknown error')}"
            
    except Exception as e:
        return f"❌ Request failed: {str(e)}"

# --- OPTIMIZATION 3: Fast Non-Streaming Fallback ---
def query_llama_agent(prompt, max_tokens=800):
    """Optimized non-streaming LLM with reduced timeout"""
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }

    payload = {
        "model": LLAMA_MODEL_ID,
        "prompt": prompt,
        "max_tokens": max_tokens,
        "temperature": 0.7,
    }

    try:
        response = requests.post(
            "https://api.together.xyz/v1/completions",
            headers=headers,
            json=payload,
            timeout=20  # Reduced from 30s
        )
        result = response.json()
        
        if response.status_code == 200 and "choices" in result:
            return result["choices"][0]["text"].strip()
        elif response.status_code == 429:
            time.sleep(1)
            return query_llama_agent(prompt, max_tokens)  # Single retry
        else:
            return f"❌ API Error: {result.get('error', 'Unknown error')}"
            
    except Exception as e:
        return f"❌ Request failed: {str(e)}"

# --- Code Execution ---
def extract_and_run_code(response_text, df, chat_id):
    """Execute Python code from LLM response"""
    code_blocks = re.findall(r"```python(.*?)```", response_text, re.DOTALL)
    
    if not code_blocks:
        return []
    
    st.subheader("🔧 Generated Code & Results")
    execution_results = []
    
    for i, code in enumerate(code_blocks):
        code = re.sub(r"df\s*=\s*pd\.read_csv\(.*?\)", "", code.strip())
        code = re.sub(r"plt\.show\(\)", "", code)
        code = code.strip()
        
        if not code:
            continue
            
        st.code(code, language="python")
        
        try:
            buffer = io.StringIO()
            old_stdout = sys.stdout
            sys.stdout = buffer
            
            plt.clf()
            plt.close('all')
            
            exec_globals = {
                "pd": pd, 
                "np": np, 
                "plt": plt, 
                "sns": sns,
                "px": px,
                "go": go,
                "df": df,
                "st": st
            }
            
            exec(code, exec_globals)
            
            sys.stdout = old_stdout
            printed_output = buffer.getvalue()
            
            result = {
                'code': code,
                'output': printed_output.strip(),
                'has_plots': bool(plt.get_fignums()),
                'execution_time': datetime.now()
            }
            execution_results.append(result)
            
            if printed_output:
                st.subheader("📊 Output")
                st.text(printed_output.strip())
            
            if plt.get_fignums():
                st.subheader("📈 Visualization")
                for fig_num in plt.get_fignums():
                    fig = plt.figure(fig_num)
                    st.pyplot(fig)
                    plt.close(fig)
            
        except Exception as e:
            sys.stdout = old_stdout
            st.error(f"❌ Execution error: {str(e)}")
            
            result = {
                'code': code,
                'error': str(e),
                'execution_time': datetime.now()
            }
            execution_results.append(result)
        finally:
            plt.close('all')
    
    return execution_results

# --- Chat History Management ---
def add_to_chat_history(question, response, data_type, execution_results=None):
    """Add conversation to history"""
    chat_entry = {
        'id': f"chat_{len(st.session_state.chat_history)}_{int(time.time())}",
        'question': question,
        'response': response,
        'timestamp': datetime.now(),
        'data_type': data_type,
        'execution_results': execution_results or [],
        'file_name': st.session_state.current_file_name
    }
    st.session_state.chat_history.append(chat_entry)
    return chat_entry['id']

def display_chat_history():
    """Display chat history"""
    if not st.session_state.chat_history:
        return
        
    st.subheader("💬 Conversation History")
    
    for i, chat in enumerate(st.session_state.chat_history):
        timestamp = chat['timestamp'].strftime("%H:%M:%S")
        
        with st.expander(f"[{timestamp}] 👤 {chat['question'][:60]}...", expanded=(i == len(st.session_state.chat_history) - 1)):
            st.markdown(f"**👤 User:** {chat['question']}")
            st.markdown(f"**🤖 Agent:**")
            st.markdown(chat['response'])
            
            if chat.get('execution_results'):
                st.markdown("**🔧 Code Execution Results:**")
                for j, result in enumerate(chat['execution_results']):
                    if 'error' in result:
                        st.error(f"Code block {j+1} failed: {result['error']}")
                    else:
                        st.success(f"Code block {j+1} executed successfully")
                        if result.get('output'):
                            st.text(result['output'])
                        if result.get('has_plots'):
                            st.info("📈 Visualization was generated")
            
            st.caption(f"File: {chat.get('file_name', 'N/A')} | Type: {chat.get('data_type', 'N/A')}")
            
            if st.button(f"🔄 Re-run this query", key=f"rerun_{chat['id']}"):
                st.session_state.selected_question = chat['question']
                st.rerun()

def export_chat_history():
    """Export chat history"""
    if not st.session_state.chat_history:
        return None
    
    export_data = []
    for chat in st.session_state.chat_history:
        export_entry = {
            'timestamp': chat['timestamp'].isoformat(),
            'file_name': chat.get('file_name', ''),
            'question': chat['question'],
            'response': chat['response'],
            'data_type': chat.get('data_type', ''),
            'execution_count': len(chat.get('execution_results', []))
        }
        export_data.append(export_entry)
    
    return pd.DataFrame(export_data)

# --- Data Insights Generator ---
def generate_data_insights(df):
    """Generate automatic insights"""
    insights = []
    insights.append(f"📊 Dataset has {df.shape[0]:,} rows and {df.shape[1]} columns")
    
    missing = df.isnull().sum()
    if missing.sum() > 0:
        insights.append(f"⚠️ Found {missing.sum():,} missing values across {(missing > 0).sum()} columns")
    
    numeric_cols = df.select_dtypes(include=[np.number]).columns
    categorical_cols = df.select_dtypes(include=['object']).columns
    
    if len(numeric_cols) > 0:
        insights.append(f"🔢 {len(numeric_cols)} numeric columns: {', '.join(numeric_cols[:3])}{'...' if len(numeric_cols) > 3 else ''}")
    
    if len(categorical_cols) > 0:
        insights.append(f"📝 {len(categorical_cols)} categorical columns: {', '.join(categorical_cols[:3])}{'...' if len(categorical_cols) > 3 else ''}")
    
    duplicate_count = df.duplicated().sum()
    if duplicate_count > 0:
        insights.append(f"🔄 Found {duplicate_count:,} duplicate rows")
    
    return insights

# --- Sidebar ---
with st.sidebar:
    st.header("🛠️ Configuration")
    
    # OPTIMIZATION: Streaming toggle
    use_streaming = st.checkbox("⚡ Streaming Response", value=True, help="Real-time response display")
    
    with st.expander("Model Settings"):
        max_tokens = st.slider("Max Tokens", 100, 2048, 800)  # Lower default
        temperature = st.slider("Temperature", 0.0, 1.0, 0.7)
    
    if st.session_state.chat_history:
        st.header("💬 Chat Management")
        
        export_df = export_chat_history()
        if export_df is not None:
            csv = export_df.to_csv(index=False)
            st.download_button(
                label="📥 Export Chat",
                data=csv,
                file_name=f"chat_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                mime="text/csv"
            )
        
        st.metric("Conversations", len(st.session_state.chat_history))
        
        if st.button("🗑️ Clear History"):
            st.session_state.chat_history = []
            st.session_state.data_summary_cache = {}
            st.rerun()
    
    st.header("⚡ Quick Actions")
    if st.session_state.current_data is not None:
        if st.button("📈 Data Summary"):
            st.session_state.selected_question = 'Provide concise data summary with key stats'
        
        if st.button("🔍 Correlations"):
            st.session_state.selected_question = 'Find top correlations and create heatmap'
        
        if st.button("📊 Dashboard"):
            st.session_state.selected_question = 'Create key visualizations showing insights'
    
    if st.session_state.chat_history:
        st.header("🕒 Recent")
        for i, chat in enumerate(reversed(st.session_state.chat_history[-5:])):
            timestamp = chat['timestamp'].strftime("%H:%M")
            preview = f"{chat['question'][:25]}..."
            if st.button(f"[{timestamp}] {preview}", key=f"recent_{i}"):
                st.session_state.selected_question = chat['question']

# --- Main Interface ---
st.title("📊 Advanced Data Analyst Agent")
st.markdown("Upload your data and ask questions to get insights, analysis, and visualizations!")

uploaded_file = st.file_uploader(
    "Upload a file", 
    type=["csv", "xlsx", "pdf", "txt", "docx", "png", "jpg", "jpeg"],
    help="Supported formats: CSV, Excel, PDF, Text, Word documents, and images"
)

if uploaded_file:
    if st.session_state.current_file_name != uploaded_file.name:
        if st.session_state.chat_history:
            if st.button("🗑️ New file detected. Clear chat history?"):
                st.session_state.chat_history = []
                st.session_state.data_summary_cache = {}
        st.session_state.current_file_name = uploaded_file.name
    
    with st.spinner("Processing file..."):
        parsed_data, data_type, metadata = parse_file(uploaded_file)
    
    if parsed_data is not None:
        st.session_state.current_data = parsed_data
        st.session_state.data_info = metadata
        
        if data_type == "dataframe":
            st.subheader("📋 Dataset Overview")
            st.dataframe(parsed_data.head(10), use_container_width=True)
            
            insights = generate_data_insights(parsed_data)
            st.subheader("🔍 Quick Insights")
            for insight in insights:
                st.info(insight)
            
        elif data_type == "text":
            st.subheader("📄 Document Preview")
            preview_text = parsed_data[:1000] + "..." if len(parsed_data) > 1000 else parsed_data
            st.text_area("Content Preview", preview_text, height=200)
        
        if st.session_state.chat_history:
            display_chat_history()
        
        st.subheader("💬 Ask Your Question")
        
        if data_type == "dataframe":
            suggested_questions = [
                "What are the main patterns in this data?",
                "Show statistical summaries for key columns",
                "Create visualizations for important relationships",
                "Find outliers or data quality issues",
                "Generate analysis with key insights"
            ]
        else:
            suggested_questions = [
                "Summarize the key points",
                "What are the main themes?",
                "Extract important facts",
                "Identify recommendations"
            ]
        
        st.write("**Suggested questions:**")
        cols = st.columns(min(3, len(suggested_questions)))
        for i, suggestion in enumerate(suggested_questions):
            col_idx = i % len(cols)
            if cols[col_idx].button(suggestion, key=f"suggest_{i}"):
                st.session_state.selected_question = suggestion
        
        question = st.text_input(
            "Type your question:",
            value=getattr(st.session_state, 'selected_question', ''),
            placeholder="e.g., What trends do you see? Create visualizations."
        )
        
        col1, col2 = st.columns([1, 4])
        ask_button = col1.button("🚀 Ask", type="primary")
        
        if ask_button and question:
            with st.spinner("🤔 Analyzing..."):
                if data_type == "dataframe":
                    # OPTIMIZATION: Use compact cached summary
                    data_summary = get_compact_data_summary(parsed_data, uploaded_file.name)
                    
                    # OPTIMIZATION: Minimal conversation context (last 2 only)
                    conversation_context = ""
                    if st.session_state.chat_history:
                        for prev in st.session_state.chat_history[-2:]:
                            conversation_context += f"\nPrev Q: {prev['question'][:80]}\n"
                    
                    # OPTIMIZATION: Compact context string
                    sample_str = "\n".join([str(row) for row in data_summary['sample'][:2]])
                    
                    context = f"""Dataset: {data_summary['shape'][0]} rows × {data_summary['shape'][1]} cols
Columns: {', '.join(data_summary['columns'][:8])}
Numeric: {', '.join(data_summary['numeric_cols'][:5])}
Sample data (first 2 rows):
{sample_str}
{conversation_context}"""
                    
                    viz_keywords = ['plot', 'chart', 'graph', 'visualiz', 'show', 'display', 'draw', 'histogram', 'scatter', 'bar', 'heatmap', 'box', 'distribution']
                    needs_visualization = any(kw in question.lower() for kw in viz_keywords)
                    
                    if needs_visualization:
                        prompt = f"""Data analyst task. Answer concisely with visualization code.

{context}

Question: {question}

Provide:
1. Brief 2-sentence analysis
2. Python visualization code (matplotlib/seaborn)
3. One key insight

Code rules:
- Use 'df' variable
- plt.figure(figsize=(10,6))
- NO plt.show()
- Add labels/titles

Be concise and direct."""
                    else:
                        prompt = f"""Data analyst task. Answer the question briefly.

{context}

Question: {question}

Provide concise analysis. Include Python code only if needed for calculations. Be direct."""

                else:
                    # Text document handling
                    conversation_context = ""
                    if st.session_state.chat_history:
                        for prev in st.session_state.chat_history[-2:]:
                            conversation_context += f"\nQ: {prev['question'][:60]}\n"
                    
                    preview = parsed_data[:1500]  # Reduced from 3000
                    prompt = f"""Document analysis. Answer concisely.

Document excerpt:
{preview}
{conversation_context}

Question: {question}

Provide brief, direct answer with specific references."""

                # OPTIMIZATION: Use streaming or fast non-streaming
                st.subheader("🧠 Analysis Results")
                if use_streaming:
                    response = query_llama_streaming(prompt, max_tokens=max_tokens)
                else:
                    response = query_llama_agent(prompt, max_tokens=max_tokens)
                    st.markdown(response)
                
                execution_results = []
                if data_type == "dataframe" and "```python" in response:
                    execution_results = extract_and_run_code(response, parsed_data, st.session_state.current_chat_id)
                
                chat_id = add_to_chat_history(question, response, data_type, execution_results)
                st.session_state.current_chat_id = chat_id
            
            if hasattr(st.session_state, 'selected_question'):
                del st.session_state.selected_question

else:
    st.info("👆 Please upload a file to begin analysis")
    
    st.subheader("📝 How to use:")
    st.markdown("""
    1. **Upload data** - CSV, Excel, PDF, Word, images
    2. **Ask questions** - Natural language queries
    3. **Get insights** - Analysis, visualizations, code
    4. **Review history** - All conversations saved
    5. **Re-run queries** - Click previous questions
    
    **Example questions:**
    - "What are the trends? Show visualizations."
    - "Show correlations with heatmap"
    - "Find outliers with box plots"
    - "Create dashboard of key metrics"
    """)

st.markdown("---")
st.markdown("Built with ❤️ using Streamlit • Optimized for Speed ⚡")